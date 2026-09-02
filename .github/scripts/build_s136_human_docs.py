from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

ROOT=Path('.')
OUT=ROOT/'Current'/'HumanReadable'
OUT.mkdir(parents=True, exist_ok=True)

SOURCES={
'00':ROOT/'Current'/'00_CURRENT_STATE.md',
'01':ROOT/'Current'/'01_HANDOVER_CORE.md',
'02':ROOT/'Current'/'02_TECHNICAL_BASELINE.md',
'03':ROOT/'Current'/'03_PROJECT_CHRONOLOGY.md',
'04':ROOT/'Current'/'04_OPEN_ISSUES_AND_NEXT_TESTS.md',
'05':ROOT/'Current'/'05_FAILED_AND_OBSOLETE_APPROACHES.md',
'06':ROOT/'Current'/'06_RECENT_WORK_S1.32-S1.36.md',
}

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_cell_shading(cell, fill='D9EAF7'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Page ')
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def setup_doc(title, subtitle):
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Inches(0.75); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8)
    styles=doc.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10.5)
    for s,size in [('Title',24),('Heading 1',17),('Heading 2',14),('Heading 3',12)]:
        styles[s].font.name='Aptos'; styles[s].font.size=Pt(size)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.size=Pt(24)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p2.add_run(subtitle); r.italic=True; r.font.size=Pt(11)
    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run('Stand: 2026-09-02 | Lethal Company V81 | Canonical candidate: S1.36').font.size=Pt(10)
    doc.add_paragraph('')
    note=doc.add_paragraph(); note.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=note.add_run('Wichtig: S1.36 ist Build-/Archiv-/Diff-verifiziert, aber noch nicht runtime-getestet. Letzter Runtime-Test: S1.34.')
    rr.bold=True
    doc.add_page_break()
    for section in doc.sections:
        add_page_number(section.footer.paragraphs[0])
    return doc

inline_pat=re.compile(r'(\*\*.*?\*\*|\`.*?\`)')
def add_inline(par, text):
    pos=0
    for m in inline_pat.finditer(text):
        if m.start()>pos: par.add_run(text[pos:m.start()])
        token=m.group(0)
        if token.startswith('**'):
            r=par.add_run(token[2:-2]); r.bold=True
        elif token.startswith('`'):
            r=par.add_run(token[1:-1]); r.font.name='Consolas'; r.font.size=Pt(9.5)
        pos=m.end()
    if pos<len(text): par.add_run(text[pos:])

def parse_table(lines, i, doc):
    rows=[]
    while i < len(lines) and lines[i].strip().startswith('|'):
        row=[c.strip() for c in lines[i].strip().strip('|').split('|')]
        rows.append(row); i+=1
    if len(rows)>=2 and all(set(c.replace(':','').replace('-','').strip())==set() for c in rows[1]):
        rows.pop(1)
    if rows:
        cols=max(len(r) for r in rows)
        table=doc.add_table(rows=len(rows), cols=cols)
        table.alignment=WD_TABLE_ALIGNMENT.CENTER
        table.style='Table Grid'
        for ri,row in enumerate(rows):
            for ci in range(cols):
                txt=row[ci] if ci<len(row) else ''
                cell=table.cell(ri,ci); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text=''; add_inline(cell.paragraphs[0], txt)
                if ri==0:
                    set_cell_shading(cell)
                    for rr in cell.paragraphs[0].runs: rr.bold=True
        set_repeat_table_header(table.rows[0]); doc.add_paragraph('')
    return i

def add_markdown(doc, text, suppress_first_h1=False):
    lines=text.splitlines(); i=0; first_h1=True
    while i<len(lines):
        st=lines[i].strip()
        if not st: i+=1; continue
        if st.startswith('```'):
            code=[]; i+=1
            while i<len(lines) and not lines[i].strip().startswith('```'):
                code.append(lines[i].rstrip()); i+=1
            if i<len(lines): i+=1
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.18); p.paragraph_format.right_indent=Inches(0.18)
            r=p.add_run('\n'.join(code)); r.font.name='Consolas'; r.font.size=Pt(9)
            pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2F2F2'); pPr.append(shd)
            continue
        if st.startswith('|') and i+1<len(lines) and lines[i+1].strip().startswith('|'):
            i=parse_table(lines,i,doc); continue
        if st.startswith('# '):
            if suppress_first_h1 and first_h1:
                first_h1=False; i+=1; continue
            first_h1=False; doc.add_heading(st[2:].strip(), level=1); i+=1; continue
        if st.startswith('## '): doc.add_heading(st[3:].strip(), level=2); i+=1; continue
        if st.startswith('### '): doc.add_heading(st[4:].strip(), level=3); i+=1; continue
        if st.startswith('> '):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.right_indent=Inches(0.15)
            r=p.add_run(st[2:]); r.italic=True; i+=1; continue
        m=re.match(r'^(\d+)\.\s+(.*)$',st)
        if m:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.18); p.paragraph_format.first_line_indent=Inches(-0.18)
            add_inline(p, m.group(1)+'. '+m.group(2)); i+=1; continue
        if st.startswith('- '):
            p=doc.add_paragraph(style='List Bullet'); add_inline(p,st[2:]); i+=1; continue
        p=doc.add_paragraph(); add_inline(p,st); i+=1

def build_full():
    doc=setup_doc('Lethal Company - Vollständige Übergabe bis S1.36','Maschinenlesbarer Projektstand in menschenlesbarer Form')
    doc.add_heading('Dokumentzweck',1)
    doc.add_paragraph('Diese Fassung bündelt den aktuellen S1.36-Projektstand, Kernregeln, technische Baseline, offene Tests, Anti-Regression-Regeln, jüngste Arbeitsphase und Chronologie. Die Markdown-/TXT-/JSON-Dateien im Repository bleiben die primären Quellen für ChatGPT.')
    for k in ['00','01','02','04','05','06','03']:
        add_markdown(doc,SOURCES[k].read_text(encoding='utf-8'),suppress_first_h1=False)
    path=OUT/'Handover-Prompt_Lethal-Company_bis_S1.36.docx'; doc.save(path); return path

def build_core():
    doc=setup_doc('Lethal Company - S1.36 Aktueller Kern','Kurzübergabe für den sofortigen Einstieg')
    for k in ['00','01','04']:
        add_markdown(doc,SOURCES[k].read_text(encoding='utf-8'))
    path=OUT/'Handover-Prompt_S1.36_Aktueller-Kern.docx'; doc.save(path); return path

def build_chronology():
    doc=setup_doc('Lethal Company - Projektchronik kompakt bis S1.36','Technische Entwicklungs- und Entscheidungsrekonstruktion')
    add_markdown(doc,SOURCES['03'].read_text(encoding='utf-8'),suppress_first_h1=True)
    path=OUT/'Lethal-Company_Projektchronik_Kompakt_bis_S1.36.docx'; doc.save(path); return path

def build_chat():
    doc=setup_doc('Lethal Company - Chat-/Arbeitschronik bis S1.36','Rekonstruierte technische Übergabe, nicht als wörtliches Chat-Transkript zu verstehen')
    doc.add_heading('Hinweis zur Rekonstruktion',1)
    doc.add_paragraph('Dieses Dokument ist eine technische Rekonstruktion der Entscheidungen, Tests, Diagnosen und Build-Schritte. Es ist kein wörtliches Transkript jeder Nachricht. Für verbindliche aktuelle Fakten gelten die Current-Dateien des Repositories und die neuesten Runtime-Logs.')
    doc.add_heading('Jüngste Arbeitsphase S1.32-S1.36',1)
    add_markdown(doc,SOURCES['06'].read_text(encoding='utf-8'),suppress_first_h1=True)
    doc.add_page_break()
    doc.add_heading('Gesamte Projektchronologie',1)
    add_markdown(doc,SOURCES['03'].read_text(encoding='utf-8'),suppress_first_h1=True)
    path=OUT/'Lethal-Company_Chatverlauf_Handover_bis_S1.36.docx'; doc.save(path); return path

for fn in [build_full,build_core,build_chronology,build_chat]:
    print(fn())
